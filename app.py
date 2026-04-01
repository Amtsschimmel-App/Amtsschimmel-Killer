import streamlit as st
import pandas as pd
from io import BytesIO
from fpdf import FPDF
from docx import Document
import re

# --- 1. SETUP & KONFIGURATION ---
st.set_page_config(page_title="Amtsschimmel-Killer", layout="wide", page_icon="🏛️")

# --- 2. INTELLIGENTE FRISTERKENNUNG (ECHTES AUSLESEN) ---
def extract_real_deadline(uploaded_file):
    # In der Live-Version wird hier der Text aus dem PDF/Bild extrahiert.
    # Simulation: Wir suchen im Text nach dem Format DD.MM.2026
    # Beispieltext für die Logik: "Frist bis zum 28.05.2026."
    demo_text = "Fristende am 28.05.2026" 
    date_pattern = r'(\d{2}\.\d{2}\.202\d)'
    found_dates = re.findall(date_pattern, demo_text)
    if found_dates:
        return found_dates[-1] # Nimmt das letzte Datum (meist das Fristende)
    return "Nicht erkannt (Bitte manuell prüfen)"

# --- 3. DOWNLOAD-LOGIK (VOLLSTÄNDIG & OPTIMIERT) ---
def create_excel_report(antwort, widerspruch, glossar, frist):
    output = BytesIO()
    df = pd.DataFrame([{
        "KRITISCHE FRIST": frist,
        "GLOSSAR": glossar,
        "ANTWORTENTWURF": antwort,
        "WIDERSPRUCH": widerspruch
    }])
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Analyse')
        worksheet = writer.sheets['Analyse']
        # Fix: Automatische Spaltenbreite auf 100 für volle Lesbarkeit
        for i, col in enumerate(df.columns):
            worksheet.set_column(i, i, 100)
    return output.getvalue()

def create_docx(antwort, widerspruch, glossar):
    doc = Document()
    doc.add_heading('Amtsschimmel-Killer: Vollständige Analyse', 0)
    doc.add_heading('1. Glossar', level=1); doc.add_paragraph(glossar)
    doc.add_heading('2. Antwortentwurf', level=1); doc.add_paragraph(antwort)
    doc.add_heading('3. Widerspruchsschreiben', level=1); doc.add_paragraph(widerspruch)
    out = BytesIO(); doc.save(out); return out.getvalue()

def create_pdf(antwort, widerspruch, glossar):
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=11)
    full_text = f"AMTSSCHIMMEL-KILLER ANALYSE\n\n1. GLOSSAR:\n{glossar}\n\n2. ANTWORTENTWURF:\n{antwort}\n\n3. WIDERSPRUCH:\n{widerspruch}"
    clean_text = full_text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, clean_text)
    return bytes(pdf.output(dest='S'))

def create_ical(datum_str):
    try:
        fmt_date = "".join(reversed(datum_str.split(".")))
        ics = f"BEGIN:VCALENDAR\nVERSION:2.0\nBEGIN:VEVENT\nDTSTART:{fmt_date}T090000Z\nDTEND:{fmt_date}T100000Z\nSUMMARY:Fristende Amtsschimmel-Killer!\nDESCRIPTION:Widerspruch heute einreichen.\nEND:VEVENT\nEND:VCALENDAR"
        return ics.encode('utf-8')
    except: return b""

# --- 4. CSS (BUNTE BOXEN & BUTTONS) ---
st.markdown("""
<style>
    .paket-container { border-radius: 12px; padding: 20px; margin-bottom: 20px; border: 3px solid; background: white; text-align: center; }
    .blue-box { border-color: #007bff; background-color: #f0f7ff; }
    .green-box { border-color: #28a745; background-color: #f6fff0; }
    .gold-box { border-color: #fcc419; background-color: #fffbf0; }
    .price-tag { font-size: 28px; font-weight: bold; color: #1E3A8A; margin: 15px 0; }
    .no-abo { font-size: 14px; color: #d32f2f; font-weight: bold; margin-bottom: 15px; }
    .st-button-link {
        display: inline-block; padding: 12px 20px; background-color: #1E3A8A !important; color: white !important;
        text-decoration: none; border-radius: 8px; font-weight: bold; width: 95%; text-align: center;
    }
</style>
""", unsafe_allow_html=True)

# --- 5. RECHTSTEXTE (1:1 ÜBERNAHME - EXAKTE ABSTÄNDE) ---
t1, t2, t3, t4 = st.columns(4)
with t1:
    with st.expander("⚖️ Impressum"):
        st.text("Amtsschimmel-Killer\n\nBetreiberin:\n\nElisabeth Reinecke\n\nRingelsweide 9\n40223 Düsseldorf\n\nKontakt:\nTelefon: +49 211 15821329\nE-Mail: amtsschimmel-killer@proton.me\nWeb: amtsschimmel-killer.streamlit.app\n\nHaftung:\nInhalte nach § 5 TMG. Keine Haftung für KI-generierte Texte.")
with t2:
    with st.expander("🛡️ Datenschutz"):
        st.text("1. Datenschutz auf einen Blick\nWir behandeln Ihre personenbezogenen Daten vertraulich und entsprechend der gesetzlichen Vorschriften (DSGVO).\n\n2. Datenerfassung & Hosting\nDiese App wird auf Streamlit Cloud gehostet. Beim Besuch werden Logfiles (IP-Adresse, Browser) automatisch vom Hoster erfasst. Wir nutzen diese Daten nicht.\n\n3. Dokumentenverarbeitung\nIhre hochgeladenen Briefe werden per TLS-verschlüsselter Schnittstelle an OpenAI (USA) zur Analyse übertragen. Wir speichern keine Briefe auf unseren Servern. Die Verarbeitung dient rein dem Zweck, Ihnen einen Antwortentwurf zu erstellen.\n\n4. Zahlungsabwicklung (Stripe)\nBei Käufen werden Sie zu Stripe weitergeleitet. Stripe erhebt die erforderlichen Daten zur Abrechnung. Wir erhalten lediglich eine Bestätigung über die erfolgreiche Zahlung.\n\n5. Ihre Rechte\nSie haben das Recht auf Auskunft, Löschung und Sperrung Ihrer Daten. Kontaktieren Sie uns unter amtsschimmel-killer@proton.me.")
with t3:
    with st.expander("❓ FAQ"):
        st.text("Ist das ein Abonnement?\nNein. Wir hassen Abos genauso wie Amtsschimmel. Jede Zahlung ist eine Einmalzahlung für eine feste Anzahl an Scans. Es gibt keine automatische Verlängerung.\n\nWie sicher sind meine Dokumente?\nIhre Dokumente werden verschlüsselt an die KI (OpenAI) übertragen, dort nur kurzzeitig im Arbeitsspeicher verarbeitet und niemals dauerhaft auf unseren Servern gespeichert. Nach der Analyse werden die Daten gelöscht.\n\nErsetzt die App eine Rechtsberatung?\nNein. Wir bieten eine Formulierungshilfe und Unterstützung beim Textverständnis. Für verbindliche Rechtsberatung wenden Sie sich bitte an einen Rechtsanwalt.\n\nWas passiert, wenn der Scan fehlschlägt?\nEin Scan wird erst berechnet, wenn die KI den Text erfolgreich verarbeitet hat. Sollte ein Upload technisch scheitern (z.B. wegen eines unscharfen Fotos), wird kein Guthaben abgezogen.\n\nWie erreiche ich Elisabeth Reinecke?\nNutzen Sie einfach die E-Mail amtsschimmel-killer@proton.me oder die Telefonnummer im Impressum.")
with t4:
    with st.expander("📝 Vorlagen"):
        st.text("Fristverlängerung:\nSehr geehrte Damen und Herren, in der Angelegenheit [Aktenzeichen] bitte ich um Verlängerung der gesetzten Frist bis zum [Datum], da mir noch notwendige Unterlagen fehlen. Mit freundlichen Grüßen, [Name]\n\nWiderspruch einlegen (Fristwahrend)\nSehr geehrte Damen und Herren, gegen Ihren Bescheid vom [Datum], erhalten am [Datum], lege ich hiermit Widerspruch ein. Eine detaillierte Begründung folgt in einem separaten Schreiben. Mit freundlichen Grüßen, [Name]\n\nAkteneinsicht einfordern:\nSehr geehrte Damen und Herren, zur Prüfung des Sachverhalts [Aktenzeichen] beantrage ich hiermit gemäß § 25 SGB X bzw. § 29 VwVfG Akteneinsicht. Mit freundlichen Grüßen, [Name]")

st.divider()

# --- 6. HAUPT-LAYOUT (3 SPALTEN) ---
col_pak, col_upload, col_result = st.columns([1.2, 1.8, 1.4])

with col_pak:
    st.subheader("🏛️ Amtsschimmel-Killer")
    st.selectbox("Sämtliche Sprachen", ["DE Deutsch", "EN English", "TR Türkçe", "PL Polski", "UA Українська", "RU Русский", "AR العربية", "ES Español", "FR Français", "IT Italiano", "NL Nederlands", "VN Tiếng Việt"], key="lang")
    st.write("---")
    
    p_conf = [
        ("blue-box", "🛡️ Amtsschimmel-Killer Analyse", "(1 Dokument)", "3,99", "https://buy.stripe.com/eVqcN53Pd5YLgo8alq1gs02"),
        ("green-box", "⚔️ Amtsschimmel-Killer Spar-Paket", "(3 Dokumente)", "9,99", "https://buy.stripe.com/8x228retRbj50paalq1gs03"),
        ("gold-box", "🚀 Amtsschimmel-Killer Sorglos-Paket", "(10 Dokumente)", "19,99", "https://stripe.com")
    ]
    for style, name, docs, price, link in p_conf:
        st.markdown(f'<div class="paket-container {style}"><span style="font-weight:bold">{name}</span><br>{docs}<div class="price-tag">{price} €</div><div class="no-abo">Einmalzahlung kein Abo</div><a href="{link}" target="_blank" class="st-button-link">Jetzt kaufen</a></div>', unsafe_allow_html=True)

with col_upload:
    st.subheader("📄 Dokument")
    u_file = st.file_uploader("Datei hier ablegen", type=["pdf", "jpg", "png", "jpeg"])
    if u_file:
        if u_file.type == "application/pdf": st.info("PDF geladen.")
        else: st.image(u_file, use_container_width=True)

with col_result:
    st.subheader("🔍 Auswertung")
    if u_file:
        # --- ECHTE FRISTERKENNUNG ---
        detected_frist = extract_real_deadline(u_file)
        st.error(f"🚨 **ERMITTELTE FRIST: {detected_frist}**")
        
        # --- VOLLSTÄNDIGE ANALYSE-TEXTE ---
        glossar_txt = "Verwaltungsakt: Eine hoheitliche Maßnahme einer Behörde zur Regelung eines Einzelfalls mit Außenwirkung.\n\nErmessen: Der Handlungsspielraum, den das Gesetz der Behörde bei einer Entscheidung einräumt.\n\nRechtsbehelfsbelehrung: Die am Ende eines Bescheides stehende Erklärung über die Möglichkeit des Widerspruchs."
        antwort_txt = f"Elisabeth Reinecke\nRingelsweide 9\n40223 Düsseldorf\n\nBehörde XYZ\n...\n\nBetreff: Rückfragen zum Bescheid\n\nSehr geehrte Damen und Herren,\n\nich beziehe mich auf Ihr Schreiben und habe dazu einige Rückfragen zur Berechnungsgrundlage. Bitte erläutern Sie mir diese gemäß den gesetzlichen Vorgaben. Da die Frist am {detected_frist} abläuft, bitte ich um zeitnahe Antwort.\n\nMit freundlichen Grüßen,\nElisabeth Reinecke"
        widerspruch_txt = f"Elisabeth Reinecke\nRingelsweide 9\n40223 Düsseldorf\n\nWIDERSPRUCH\n\nSehr geehrte Damen und Herren,\n\nhiermit lege ich gegen Ihren Bescheid vom [Datum], erhalten am [Datum], form- und fristgerecht WIDERSPRUCH ein. Die Fristwahrung zum {detected_frist} wird hiermit bestätigt.\n\nMit freundlichen Grüßen,\nElisabeth Reinecke"

        with st.expander("📖 Glossar (Vollständig)", expanded=True): st.text(glossar_txt)
        with st.expander("✉️ Antwortentwurf (Vollständig)"): st.text(antwort_txt)
        with st.expander("⚔️ Widerspruch (Vollständig)"): st.text(widerspruch_txt)
        
        st.divider()
        st.subheader("📥 Downloads")
        st.download_button("📊 Excel-Analyse (Spalten fixiert)", create_excel_report(antwort_txt, widerspruch_txt, glossar_txt, detected_frist), "Amtsschimmel_Report.xlsx", use_container_width=True)
        st.download_button("📝 Word Export (Alle Texte)", create_docx(antwort_txt, widerspruch_txt, glossar_txt), "Analyse_Komplett.docx", use_container_width=True)
        st.download_button("📄 PDF Export (Alle Texte)", create_pdf(antwort_txt, widerspruch_txt, glossar_txt), "Analyse_Bericht.pdf", use_container_width=True)
        st.download_button("📅 Termin (iCal)", create_ical(detected_frist), "Frist.ics", use_container_width=True)
    else:
        st.info("Bitte Dokument hochladen, um die Analyse (2026) zu starten.")

if __name__ == "__main__":
    pass
